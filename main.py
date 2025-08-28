import os
import io
import json
import re
from typing import Dict, Any

# Step 0: Import necessary libraries
# azure-ai-formrecognizer: For PDF text extraction and OCR
# azure-openai: For structuring extracted text into JSON


from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image
import fitz  # PyMuPDF
from dotenv import load_dotenv

# Step 1: Load environment variables from .env file
# This loads Azure keys for secure access without hardcoding
load_dotenv()

# Step 2: Define constants
# EMU_PER_MM: Conversion factor for Word's internal units (EMU) from millimeters
EMU_PER_MM = 36000  # Word EMU units per mm

# Step 3: Helper function to convert mm to EMU
# Used for positioning elements in DOCX
def mm_to_emu(mm: float) -> int:
    return int(mm * EMU_PER_MM)

# Step 4: Extract text from PDF using Azure Form Recognizer
# This handles OCR for image-based PDFs, extracting lines in order
def fr_extract_text(pdf_bytes: bytes) -> str:
    client = DocumentAnalysisClient(
        endpoint=os.environ["AZURE_FORMREC_ENDPOINT"],
        credential=AzureKeyCredential(os.environ["AZURE_FORMREC_KEY"])
    )
    poller = client.begin_analyze_document(
        model_id="prebuilt-layout",
        document=pdf_bytes
    )
    result = poller.result()
    chunks = []
    for page in result.pages:
        for line in page.lines or []:
            chunks.append(line.content or "")
        chunks.append("\n")
    return "\n".join(chunks).strip()

# Step 5: Structure extracted text into JSON using Azure OpenAI
# Prompts the model to parse and formalize resume data
def aoai_structured_json(raw_text: str) -> Dict[str, Any]:
    client = AzureOpenAI(
        azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
        api_key=os.environ["AZURE_OPENAI_KEY"],
        api_version="2024-02-15-preview"  # Adjust if needed
    )
    system = "You are an expert resume parser. Output STRICT JSON only. No markdown."
    user = f"""
Extract the following fields from resume text; rephrase narrative into formal corporate tone:
JSON schema:
{{
  "name": "string",
  "contact": {{"phone":"string","email":"string","linkedin":"string"}},
  "summary": "string",
  "experience": [{{"title":"string","company":"string","dates":"string","description":["string"]}}],
  "education": [{{"degree":"string","institution":"string","dates":"string"}}],
  "skills": ["string"],
  "certifications": ["string"],
  "achievements": ["string"]
}}
Use "" or [] when unknown. Resume text:
{raw_text}
"""
    resp = client.chat.completions.create(
        model=os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o"),
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        temperature=0
    )
    content = resp.choices[0].message.content.strip()
    # Robust JSON parse to handle any extra text
    start = content.find("{")
    if start == -1:
        return {}
    depth = 0
    for i, ch in enumerate(content[start:], start):
        if ch == "{": depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                try:
                    return json.loads(content[start:i+1])
                except:
                    break
    return {}

# Step 6: Extract first image from PDF (for logo or photo)
# Uses PyMuPDF to find and return the first embedded image
def extract_first_image_from_pdf(pdf_bytes: bytes) -> Image.Image | None:
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            for img in page.get_images(full=True):
                xref = img[0]
                base_img = doc.extract_image(xref)
                ib = base_img["image"]
                if ib:
                    return Image.open(io.BytesIO(ib)).convert("RGB")
    return None

# Step 7: Add absolutely positioned textbox to DOCX using OXML
# Creates an anchored drawing with a textbox at specific coordinates
def add_textbox_anchor(paragraph, x_emu, y_emu, w_emu, h_emu, text, font_size_pt=11, bold=False):
    r = paragraph.add_run()
    drawing = OxmlElement('w:drawing')
    r._r.append(drawing)

    anchor = OxmlElement('wp:anchor')
    anchor.set(qn('distT'), "0"); anchor.set(qn('distB'), "0"); anchor.set(qn('distL'), "0"); anchor.set(qn('distR'), "0")
    anchor.set(qn('simplePos'), "0"); anchor.set(qn('relativeHeight'), "251658240")
    anchor.set(qn('behindDoc'), "0"); anchor.set(qn('locked'), "0"); anchor.set(qn('layoutInCell'), "1"); anchor.set(qn('allowOverlap'), "1")

    simple_pos = OxmlElement('wp:simplePos'); simple_pos.set(qn('x'), "0"); simple_pos.set(qn('y'), "0")
    anchor.append(simple_pos)

    posh = OxmlElement('wp:positionH'); posh.set(qn('relativeFrom'), "page")
    poshoff = OxmlElement('wp:posOffset'); poshoff.text = str(x_emu); posh.append(poshoff)
    anchor.append(posh)

    posv = OxmlElement('wp:positionV'); posv.set(qn('relativeFrom'), "page")
    posvoff = OxmlElement('wp:posOffset'); posvoff.text = str(y_emu); posv.append(posvoff)
    anchor.append(posv)

    extent = OxmlElement('wp:extent'); extent.set(qn('cx'), str(w_emu)); extent.set(qn('cy'), str(h_emu))
    anchor.append(extent)

    effect = OxmlElement('wp:effectExtent')
    effect.set(qn('l'), "0"); effect.set(qn('t'), "0"); effect.set(qn('r'), "0"); effect.set(qn('b'), "0")
    anchor.append(effect)

    graphic = OxmlElement('a:graphic')
    graphic_data = OxmlElement('a:graphicData'); graphic_data.set(qn('uri'), "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingShape")
    sp = OxmlElement('wps:sp')
    txbx = OxmlElement('wps:txbx')
    txbx_content = OxmlElement('w:txbxContent')

    p = OxmlElement('w:p'); r2 = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = text
    rpr = OxmlElement('w:rPr')
    if bold:
        b = OxmlElement('w:b'); rpr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(font_size_pt * 2)))
    rpr.append(sz); r2.append(rpr); r2.append(t); p.append(r2)
    txbx_content.append(p); txbx.append(txbx_content); sp.append(txbx); graphic_data.append(sp)
    graphic.append(graphic_data); anchor.append(graphic)
    drawing.append(anchor)

# Step 8: Add absolutely positioned image to DOCX using OXML
# Embeds an image at specific coordinates (for logos or photos)
def add_image_anchor(paragraph, x_emu, y_emu, w_emu, h_emu, image_bytes: bytes):
    part = paragraph.part
    image_part = part.add_image_part('image/png')
    image_part.blob = image_bytes
    r_id = part.relate_to(image_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

    r = paragraph.add_run()
    drawing = OxmlElement('w:drawing')
    r._r.append(drawing)

    anchor = OxmlElement('wp:anchor')
    anchor.set(qn('distT'), "0"); anchor.set(qn('distB'), "0"); anchor.set(qn('distL'), "0"); anchor.set(qn('distR'), "0")
    anchor.set(qn('simplePos'), "0"); anchor.set(qn('relativeHeight'), "251658240")
    anchor.set(qn('behindDoc'), "0"); anchor.set(qn('locked'), "0"); anchor.set(qn('layoutInCell'), "1"); anchor.set(qn('allowOverlap'), "1")

    simple_pos = OxmlElement('wp:simplePos'); simple_pos.set(qn('x'), "0"); simple_pos.set(qn('y'), "0")
    anchor.append(simple_pos)

    posh = OxmlElement('wp:positionH'); posh.set(qn('relativeFrom'), "page")
    poshoff = OxmlElement('wp:posOffset'); poshoff.text = str(x_emu); posh.append(poshoff)
    anchor.append(posh)

    posv = OxmlElement('wp:positionV'); posv.set(qn('relativeFrom'), "page")
    posvoff = OxmlElement('wp:posOffset'); posvoff.text = str(y_emu); posv.append(posvoff)
    anchor.append(posv)

    extent = OxmlElement('wp:extent'); extent.set(qn('cx'), str(w_emu)); extent.set(qn('cy'), str(h_emu))
    anchor.append(extent)

    effect = OxmlElement('wp:effectExtent')
    effect.set(qn('l'), "0"); effect.set(qn('t'), "0"); effect.set(qn('r'), "0"); effect.set(qn('b'), "0")
    anchor.append(effect)

    wrap_none = OxmlElement('wp:wrapNone')
    anchor.append(wrap_none)

    doc_elem = OxmlElement('wp:docPr')
    doc_elem.set(qn('id'), "1"); doc_elem.set(qn('name'), "Picture 1")
    anchor.append(doc_elem)

    graphic = OxmlElement('a:graphic')
    graphic_data = OxmlElement('a:graphicData'); graphic_data.set(qn('uri'), "http://schemas.openxmlformats.org/drawingml/2006/picture")
    pic = OxmlElement('pic:pic')

    nv_pic_pr = OxmlElement('pic:nvPicPr')
    c_nv_pr = OxmlElement('pic:cNvPr'); c_nv_pr.set(qn('id'), "0"); c_nv_pr.set(qn('name'), "Picture")
    nv_pic_pr.append(c_nv_pr)
    pic.append(nv_pic_pr)

    blip_fill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip'); blip.set(qn('r:embed'), r_id)
    blip_fill.append(blip)
    stretch = OxmlElement('a:stretch'); fill_rect = OxmlElement('a:fillRect'); stretch.append(fill_rect); blip_fill.append(stretch)
    pic.append(blip_fill)

    sp_pr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off'); off.set(qn('x'), "0"); off.set(qn('y'), "0")
    ext = OxmlElement('a:ext'); ext.set(qn('cx'), str(w_emu)); ext.set(qn('cy'), str(h_emu))
    xfrm.append(off); xfrm.append(ext)
    sp_pr.append(xfrm)
    pic.append(sp_pr)

    graphic_data.append(pic)
    graphic.append(graphic_data)
    anchor.append(graphic)
    drawing.append(anchor)

# Step 9: Main processing function
# Loads files, extracts, structures, and fills the template with positioned elements
def process_resume(template_docx_path: str, coords_json_path: str, resume_path: str, out_path: str):
    coords = json.load(open(coords_json_path, "r", encoding="utf-8"))
    with open(resume_path, "rb") as f:
        resume_bytes = f.read()

    raw_text = fr_extract_text(resume_bytes)
    data = aoai_structured_json(raw_text)

    if not data.get("name"):
        data["name"] = "Candidate"

    doc = Document(template_docx_path)

    anchor_host = doc.add_paragraph()  # Hidden anchor host

    def put_text(key, text, style=None, font_size=None, bold=False):
        cfg = coords["fields"].get(key)
        if not cfg or not text: return
        x = mm_to_emu(cfg["x_mm"]); y = mm_to_emu(cfg["y_mm"])
        w = mm_to_emu(cfg["w_mm"]); h = mm_to_emu(cfg["h_mm"])
        font_cfg = cfg.get("font", {})
        add_textbox_anchor(
            anchor_host, x, y, w, h, text,
            font_size_pt=font_cfg.get("size_pt", font_size or 11),
            bold=font_cfg.get("bold", bold)
        )

    # Simple fields
    put_text("NAME", data.get("name", "").upper(), bold=True)
    c = data.get("contact", {})
    put_text("PHONE", c.get("phone", ""))
    put_text("EMAIL", c.get("email", ""))
    put_text("LINKEDIN", c.get("linkedin", ""))
    put_text("SUMMARY", data.get("summary", ""))

    # List fields as \n joined text
    def lines_for_education(edus):
        return [" | ".join([b for b in [e.get("degree", ""), e.get("institution", ""), e.get("dates", "")] if b]) for e in edus]

    def lines_for_experience(exps):
        out = []
        for ex in exps:
            header = " | ".join([b for b in [ex.get("title", ""), ex.get("company", ""), ex.get("dates", "")] if b])
            if header: out.append(header)
            out += [f"• {d.strip()}" for d in ex.get("description", []) if d.strip()]
            out.append("")  # spacer
        return out

    def place_list(key, lines):
        if not lines: return
        cfg = coords["fields"].get(key); if not cfg: return
        text_block = "\n".join(lines)
        put_text(key, text_block)

    place_list("EDUCATION", lines_for_education(data.get("education", [])))
    place_list("SKILLS", data.get("skills", []))
    place_list("CERTS", data.get("certifications", []))
    place_list("EXPERIENCE", lines_for_experience(data.get("experience", [])))
    place_list("ACHIEVEMENTS", data.get("achievements", []))

    # Optional logo
    if "LOGO_SLOT" in coords["fields"]:
        img = extract_first_image_from_pdf(resume_bytes)
        if img:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            image_bytes = buf.getvalue()
            cfg = coords["fields"]["LOGO_SLOT"]
            x = mm_to_emu(cfg["x_mm"]); y = mm_to_emu(cfg["y_mm"])
            w = mm_to_emu(cfg["w_mm"]); h = mm_to_emu(cfg["h_mm"])
            add_image_anchor(anchor_host, x, y, w, h, image_bytes)

    doc.save(out_path)
    print(f"Filled resume saved to {out_path}")

# Step 10: Run the project if executed as script
# Calls the main function with predefined paths
if __name__ == "__main__":
    process_resume(
        template_docx_path="template/Resume-template-use.docx",
        coords_json_path="coords.json",
        resume_path="input/shivam-cv.pdf",
        out_path="output/filled-resume.docx"
    )
