import os
import io
import re
import time
import json
import hashlib
import base64
from collections import Counter, defaultdict
from typing import Dict, Any, List, Tuple

import fitz  # PyMuPDF
from PIL import Image
import google.generativeai as genai
import streamlit as st

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =========================
# Setup
# =========================
# Configure page
st.set_page_config(
    page_title="Formal Resume Standardizer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize Gemini (will ask for API key in sidebar)
def init_gemini():
    """Initialize Gemini with API key from user input"""
    with st.sidebar:
        st.header("API Configuration")
        api_key = st.text_input("Enter Google Gemini API Key:", type="password")
        
        if api_key:
            try:
                genai.configure(api_key=api_key)
                st.success("API key configured successfully!")
                return True
            except Exception as e:
                st.error(f"Error configuring API: {str(e)}")
                return False
        else:
            st.info("Please enter your Gemini API key to continue")
            return False

# =========================
# Constants
# =========================
OCR_MODEL = "gemini-1.5-flash"
EXTRACT_MODEL = "gemini-1.5-flash"
MAX_OCR_IMAGES = 10
OCR_DELAY_SEC = 0.1
MIN_TEXT_LEN_TO_SKIP_OCR = 100

# =========================
# Retry helper
# =========================
def retry(times=3, base_delay=0.7, factor=2.0, exceptions=(Exception,)):
    def decorator(fn):
        def wrapper(*args, **kwargs):
            delay = base_delay
            last = None
            for i in range(times):
                try:
                    return fn(*args, **kwargs)
                except exceptions as e:
                    last = e
                    if i == times - 1:
                        raise
                    time.sleep(delay)
                    delay *= factor
            raise last
        return wrapper
    return decorator

# =========================
# OCR + Parsing Helpers
# =========================
@retry(times=3, base_delay=0.8)
def _gemini_generate(model_name: str, parts: list):
    model = genai.GenerativeModel(model_name)
    return model.generate_content(parts)

_OCR_CACHE: Dict[str, str] = {}

def _img_hash(pil_img: Image.Image) -> str:
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG")
    return hashlib.sha256(buf.getvalue()).hexdigest()

def ocr_with_gemini(pil_img, model_name=OCR_MODEL) -> str:
    """Return plain text OCR via Gemini Vision with caching."""
    try:
        h = _img_hash(pil_img)
    except Exception:
        h = None
        
    if h and h in _OCR_CACHE:
        return _OCR_CACHE[h]

    prompt = (
        "Extract ONLY the legible text visible in this image.\n"
        "Preserve line breaks and order.\n"
        "Return plain text only."
    )
    
    with st.spinner("Performing OCR on image..."):
        resp = _gemini_generate(model_name, [prompt, pil_img])
    
    text = (getattr(resp, "text", "") or "").strip()
    if h:
        _OCR_CACHE[h] = text
    return text

def _coerce_xref(image_field):
    if image_field is None:
        return None
    if isinstance(image_field, int):
        return image_field
    if isinstance(image_field, dict):
        for k in ("xref", "number", "id"):
            v = image_field.get(k)
            if isinstance(v, int):
                return v
            if isinstance(v, (bytes, bytearray)):
                try: return int(v.decode("ascii"))
                except Exception: pass
            if isinstance(v, str) and v.isdigit(): return int(v)
        return None
    if isinstance(image_field, (bytes, bytearray)):
        try: return int(image_field.decode("ascii"))
        except Exception: return None
    if isinstance(image_field, str) and image_field.isdigit():
        return int(image_field)
    return None

def _pil_from_bbox(page, bbox, scale=2.0):
    rect = fitz.Rect(bbox)
    matrix = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=matrix, clip=rect, alpha=False)
    mode = "RGB" if pix.n < 4 else "RGBA"
    img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
    if img.mode != "RGB":
        img = img.convert("RGB")
    return img

def extract_linear_text_from_pdf(pdf_bytes: bytes):
    """Extract text and OCR image blocks in reading order (only when needed)."""
    linear_text = []
    ocr_count = 0
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf:
        total_pages = len(pdf)
        
        for page_num, page in enumerate(pdf, start=1):
            status_text.text(f"Processing page {page_num} of {total_pages}...")
            progress_bar.progress(page_num / total_pages)
            
            layout = page.get_text("dict")
            blocks = layout.get("blocks", [])

            # First pass: calculate total text length on page
            page_text_len = 0
            for block in blocks:
                if block.get("type", 0) == 0:  # text block
                    text_parts = []
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        if spans:
                            text_parts.append("".join(span.get("text", "") for span in spans))
                    text = "\n".join([t for t in text_parts if t is not None]).strip()
                    page_text_len += len(text)

            do_ocr = page_text_len < MIN_TEXT_LEN_TO_SKIP_OCR

            # Second pass: extract in order
            for block in blocks:
                if block.get("type", 0) == 0:  # text block
                    text_parts = []
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        if spans:
                            text_parts.append("".join(span.get("text", "") for span in spans))
                    text = "\n".join([t for t in text_parts if t is not None]).strip()
                    if text:
                        linear_text.append(text)
                elif block.get("type", 0) == 1:  # image block
                    if not do_ocr or (MAX_OCR_IMAGES is not None and ocr_count >= MAX_OCR_IMAGES):
                        continue

                    bbox = block.get("bbox", None)
                    image_field = block.get("image", None)
                    xref = _coerce_xref(image_field)
                    pil_img = None
                    
                    try:
                        if xref is not None:
                            base_img = pdf.extract_image(xref)
                            image_bytes = base_img.get("image")
                            if image_bytes:
                                pil_img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                    except Exception:
                        pass
                        
                    if pil_img is None and bbox is not None:
                        try:
                            pil_img = _pil_from_bbox(page, bbox, scale=2.0)
                        except Exception:
                            continue
                            
                    if pil_img is None:
                        continue

                    try:
                        ocr_text = ocr_with_gemini(pil_img, OCR_MODEL)
                    except Exception:
                        ocr_text = ""
                        
                    if ocr_text:
                        linear_text.append(ocr_text)
                    ocr_count += 1
                    if OCR_DELAY_SEC:
                        time.sleep(OCR_DELAY_SEC)
    
    progress_bar.empty()
    status_text.empty()
    extracted_text = "\n\n".join([t for t in linear_text if t.strip()])
    return extracted_text

def extract_text_from_docx(docx_bytes: bytes):
    doc = Document(io.BytesIO(docx_bytes))
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return text

# =========================
# Keyword/Skill Extraction
# =========================
STOPWORDS = set("""
a an the and or but if while for to from in on at by of with as into onto over under
is are was were be been being can could should would may might must do does did done
this that these those it its your you we our ours they their them i me my mine
per via etc et al using used use highly result results responsibility responsibilities
worked working work professional experienced experience role roles summary objective
""".split())

DOMAIN_HINTS = set("""
python java javascript typescript react angular vue node express django flask fastapi
spring hibernate dotnet csharp c++ go golang ruby rails php laravel kotlin swift
sql mysql postgresql sqlite mongodb redis kafka rabbitmq spark hadoop hive airflow
aws azure gcp docker kubernetes k8s terraform ansible git github gitlab ci cd cicd
ml ai nlp deep-learning machine-learning pytorch tensorflow keras scikit-learn
powerbi tableau excel snowflake databricks bigquery redshift looker
jira confluence agile scrum kanban microservices rest graphql grpc api
linux windows macos bash shell powershell cybersecurity networking devops jenkins
""".split())

def tokenize(text: str):
    text = re.sub(r"[^\w\-\+#/.\s]", " ", text, flags=re.UNICODE)
    return [t for t in re.split(r"\s+", text) if t]

def extract_keywords(text: str, top_k: int = 50):
    if not text:
        return []
    raw = tokenize(text)
    low = [t.lower() for t in raw if len(t) > 1]
    unigrams = [
        t for t in low
        if not t.isdigit() and t not in STOPWORDS and re.search(r"[a-zA-Z0-9]", t)
    ]
    u_counts = Counter(unigrams)
    bigrams = []
    for i in range(len(low) - 1):
        a, b = low[i], low[i + 1]
        if a not in STOPWORDS and b not in STOPWORDS and not (a.isdigit() or b.isdigit()):
            bigrams.append(f"{a} {b}")
    b_counts = Counter(bigrams)

    scores = defaultdict(float)
    for w, c in u_counts.items():
        s = float(c)
        if w in DOMAIN_HINTS:
            s *= 1.8
        scores[w] += s
    for w, c in b_counts.items():
        scores[w] += 1.5 * c

    ranked = sorted(scores.items(), key=lambda x: (-x[1], x[0]))
    keys = [w for w, _ in ranked]

    stylers = {"c#":"C#","c++":"C++","dotnet":".NET","k8s":"Kubernetes",
               "aws":"AWS","gcp":"GCP","sql":"SQL","api":"API","ai":"AI","ml":"ML","nlp":"NLP"}
    def pretty(k):
        if " " in k:
            return " ".join([s.upper() if s in {"aws","gcp","sql","api","ai","ml","nlp"} else s.capitalize()
                             for s in k.split()])
        return stylers.get(k, k.capitalize() if k.isalpha() else k)

    seen, out = set(), []
    for k in keys:
        pk = pretty(k)
        if pk.lower() not in seen:
            out.append(pk)
            seen.add(pk.lower())
        if len(out) >= top_k:
            break
    return out

# =========================
# LLM Extraction - FORMAL FORMAT ONLY
# =========================
def _clean_json_fence(s: str) -> str:
    s = s.strip()
    if s.startswith("```"):
        first_newline = s.find("\n")
        if first_newline != -1:
            s = s[first_newline + 1 :]
    if s.endswith("```"):
        s = s[:-3]
    return s.strip()

@retry(times=3, base_delay=0.8)
def extract_and_structure_resume(resume_text) -> Dict[str, Any]:
    model = genai.GenerativeModel(EXTRACT_MODEL)
    prompt = f"""
You are an expert resume parser. From the following resume text, extract all relevant information including:
- Name
- Contact details (phone, email, linkedin)
- Professional summary
- Work experience (job titles, companies, dates, descriptions as bullet points)
- Education (degrees, institutions, dates)
- Skills (list)
- Certifications (list)
- Achievements (list)

Rephrase ALL narrative parts to be FORMAL, PROFESSIONAL, and BUSINESS-APPROPRIATE.
Use formal business language, professional terminology, and corporate-style phrasing.

Output strictly as JSON:
{{
  "name": "string",
  "contact": {{"phone": "string","email": "string","linkedin": "string"}},
  "summary": "string",
  "experience": [{{"title": "string","company": "string","dates": "string","description": ["string"]}}],
  "education": [{{"degree": "string","institution": "string","dates": "string"}}],
  "skills": ["string"],
  "certifications": ["string"],
  "achievements": ["string"]
}}

If information is missing, use empty string or empty list.
Resume text:
{resume_text}
"""
    with st.spinner("Extracting and structuring resume data with AI..."):
        resp = model.generate_content(prompt)
    
    json_str = _clean_json_fence(resp.text or "")
    data = json.loads(json_str)
    
    # Ensure keys
    data.setdefault("contact", {})
    data.setdefault("experience", [])
    data.setdefault("education", [])
    data.setdefault("skills", [])
    data.setdefault("certifications", [])
    data.setdefault("achievements", [])
    
    return data

def guess_name_from_text(text: str) -> str:
    for line in (text or "").splitlines():
        s = line.strip()
        if 3 <= len(s) <= 60 and not any(ch.isdigit() for ch in s):
            words = s.split()
            if 1 <= len(words) <= 4 and all(re.match(r"^[A-Za-z\.\-]+$", w) for w in words):
                name = " ".join(w.capitalize() for w in words)
                return name
    return ""

def merge_skills(structured: Dict[str, Any], mined: List[str], max_items=50) -> List[str]:
    skills = [s.strip() for s in structured.get("skills", []) if isinstance(s, str) and s.strip()]
    seen, out = set(), []
    def add(x):
        k = x.strip().lower()
        if k and k not in seen:
            out.append(x.strip())
            seen.add(k)
    for s in skills: add(s)
    for k in mined: add(k)
    return out[:max_items]

# =========================
# Standard Template Generation - FORMAL FORMAT
# =========================
def create_formal_resume_docx(data: Dict[str, Any]) -> Document:
    """Create a formal standardized resume document based on extracted data with enhanced looks and colors"""
    doc = Document()
    
    # Set default font for the document
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    # Define enhanced styles with colors
    title_style = doc.styles.add_style('FormalTitle', 1)
    title_style.base_style = doc.styles['Normal']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 32, 96)  # Dark navy blue for name
    
    heading_style = doc.styles.add_style('FormalHeading', 1)
    heading_style.base_style = doc.styles['Normal']
    heading_style.font.size = Pt(12)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 51, 102)  # Professional blue for headings
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    subheading_style = doc.styles.add_style('FormalSubheading', 1)
    subheading_style.base_style = doc.styles['Normal']
    subheading_style.font.size = Pt(11)
    subheading_style.font.bold = True
    subheading_style.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray for subheadings
    subheading_style.paragraph_format.space_after = Pt(3)
    
    normal_style = doc.styles.add_style('FormalNormal', 1)
    normal_style.base_style = doc.styles['Normal']
    normal_style.font.color.rgb = RGBColor(0, 0, 0)  # Black for body text
    
    bullet_style = doc.styles.add_style('FormalBullet', 1)
    bullet_style.base_style = doc.styles['Normal']
    bullet_style.paragraph_format.left_indent = Pt(18)
    bullet_style.paragraph_format.space_after = Pt(3)
    
    # Add name as title
    name = data.get('name', '')
    if name:
        p = doc.add_paragraph(name.upper())
        p.style = 'FormalTitle'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
    
    # Add contact information
    contact = data.get('contact', {})
    contact_parts = []
    if contact.get('phone'): contact_parts.append(f"Phone: {contact['phone']}")
    if contact.get('email'): contact_parts.append(f"Email: {contact['email']}")
    if contact.get('linkedin'): contact_parts.append(f"LinkedIn: {contact['linkedin']}")
    
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.style = 'FormalNormal'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
    
    # Helper function to add horizontal line
    def add_horizontal_line():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        
        pPr = p._p.pPr
        if pPr is None:
            pPr = p._p.add_pPr()
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '003366')  # Hex for RGB(0,51,102)
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    # Add professional summary
    summary = data.get('summary', '')
    if summary:
        doc.add_paragraph("PROFESSIONAL SUMMARY", style='FormalHeading')
        p = doc.add_paragraph(summary)
        p.style = 'FormalNormal'
        add_horizontal_line()
    
    # Add professional experience
    experience = data.get('experience', [])
    if experience:
        doc.add_paragraph("PROFESSIONAL EXPERIENCE", style='FormalHeading')
        for exp in experience:
            title = exp.get('title', '').title()
            company = exp.get('company', '').title()
            dates = exp.get('dates', '')
            
            # Add position header
            header_text = f"{title}"
            if company:
                header_text += f" | {company}"
            if dates:
                header_text += f" | {dates}"
                
            p = doc.add_paragraph(header_text)
            p.style = 'FormalSubheading'
            
            # Add description bullets
            description = exp.get('description', [])
            for desc in description:
                if desc.strip():
                    p = doc.add_paragraph(f"• {desc.strip()}")
                    p.style = 'FormalBullet'
            
            p.paragraph_format.space_after = Pt(6)  # Space after last bullet
        
        add_horizontal_line()
    
    # Add education
    education = data.get('education', [])
    if education:
        doc.add_paragraph("EDUCATION", style='FormalHeading')
        for edu in education:
            degree = edu.get('degree', '').title()
            institution = edu.get('institution', '').title()
            dates = edu.get('dates', '')
            
            edu_text = f"{degree}"
            if institution:
                edu_text += f", {institution}"
            if dates:
                edu_text += f" | {dates}"
                
            p = doc.add_paragraph(f"• {edu_text}")
            p.style = 'FormalBullet'
        
        add_horizontal_line()
    
    # Add technical skills
    skills = data.get('skills', [])
    if skills:
        doc.add_paragraph("TECHNICAL SKILLS", style='FormalHeading')
        skills_text = " • ".join([s.title() for s in skills])
        p = doc.add_paragraph(skills_text)
        p.style = 'FormalNormal'
        add_horizontal_line()
    
    # Add professional certifications
    certifications = data.get('certifications', [])
    if certifications:
        doc.add_paragraph("PROFESSIONAL CERTIFICATIONS", style='FormalHeading')
        for cert in certifications:
            p = doc.add_paragraph(f"• {cert.title()}")
            p.style = 'FormalBullet'
        
        add_horizontal_line()
    
    # Add professional achievements
    achievements = data.get('achievements', [])
    if achievements:
        doc.add_paragraph("PROFESSIONAL ACHIEVEMENTS", style='FormalHeading')
        for achievement in achievements:
            p = doc.add_paragraph(f"• {achievement}")
            p.style = 'FormalBullet'
    
    return doc

# =========================
# Streamlit UI - FORMAL ONLY VERSION
# =========================
def main():
    st.title("📄 Formal Resume Standardizer")
    st.markdown("""
    **Professional Resume Conversion Tool**  
    Convert your resume to a formal, professional format suitable for corporate applications.
    """)
    
    # Initialize Gemini
    if not init_gemini():
        st.stop()
    
    # File upload section
    st.subheader("Upload Your Resume")
    uploaded_resume = st.file_uploader(
        "Choose a PDF or DOCX file", 
        type=["pdf", "docx"],
        help="Upload your resume in PDF or DOCX format for formal conversion"
    )
    
    # Fixed settings for formal format
    max_keywords = st.slider(
        "Maximum skills to include:",
        min_value=15,
        max_value=40,
        value=25,
        help="Select the number of key skills to highlight in your formal resume"
    )
    
    if uploaded_resume is not None and st.button("Convert to Formal Resume", type="primary"):
        # Extract file extension
        resume_ext = uploaded_resume.name.split('.')[-1].lower()
        resume_bytes = uploaded_resume.getvalue()
        
        # Display file info
        st.info(f"Processing {uploaded_resume.name} for formal conversion")
        
        # Extract text from resume
        with st.spinner("Extracting text from your resume..."):
            if resume_ext == "pdf":
                resume_text = extract_linear_text_from_pdf(resume_bytes)
            elif resume_ext == "docx":
                resume_text = extract_text_from_docx(resume_bytes)
            else:
                st.error("Unsupported file format")
                st.stop()
        
        if not resume_text.strip():
            st.error("No text could be extracted from the resume. Please try another file.")
            st.stop()
        
        # Structure data with AI (formal format only)
        structured_data = extract_and_structure_resume(resume_text)
        
        # Extract and merge keywords
        with st.spinner("Extracting professional skills..."):
            extracted_keywords = extract_keywords(resume_text, top_k=max_keywords*2)
            structured_data["skills"] = merge_skills(
                structured_data, extracted_keywords, max_items=max_keywords
            )
        
        # Create formal standardized resume
        with st.spinner("Creating formal professional resume..."):
            doc = create_formal_resume_docx(structured_data)
            
            # Save to bytes
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
        
        # Display success and download option
        st.success("✅ Formal resume created successfully!")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Preview structured data
            with st.expander("Preview Extracted Information"):
                st.json(structured_data)
        
        with col2:
            # Download button
            st.download_button(
                label="📥 Download Formal Resume",
                data=bio,
                file_name="formal_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Show professional stats
            st.metric("Professional Skills", len(structured_data.get("skills", [])))
            st.metric("Work Experience Items", len(structured_data.get("experience", [])))
            st.metric("Education Credentials", len(structured_data.get("education", [])))

if __name__ == "__main__":
    main()